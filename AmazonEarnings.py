from openai import OpenAI
import os
from docx import Document
from pydub import AudioSegment


client = OpenAI(
    # defaults to os.environ.get("OPENAI_API_KEY")
    # api_key="My API Key",
)


def transcribe_audio(audio_file_path):
    audio = AudioSegment.from_wav(audio_file_path)
    chunk_size = 1 * 60 * 1000  # 20 minutes in milliseconds
    chunks = [audio[i:i+chunk_size] for i in range(0, len(audio), chunk_size)]

    transcriptions = []
    for i, chunk in enumerate(chunks):
        chunk_file_path = f"temp_chunk_{i}.wav"
        chunk.export(chunk_file_path, format="wav")
        with open(chunk_file_path, 'rb') as audio_file:
            transcription = client.audio.transcriptions.create(
                file=audio_file,
                model="whisper-1"
            )
        transcriptions.append(transcription.text)
        os.remove(chunk_file_path)  # Remove the temporary chunk file

    full_transcription = ' '.join(transcriptions)
    return full_transcription

def meeting_minutes(transcription):
    abstract_summary = abstract_summary_extraction(transcription)
    key_points = key_points_extraction(transcription)
    action_items = action_item_extraction(transcription)
    sentiment = sentiment_analysis(transcription)
    return {
        'abstract_summary': abstract_summary,
        'key_points': key_points,
        'action_items': action_items,
        'sentiment': sentiment
    }

def abstract_summary_extraction(transcription):
    response = client.chat.completions.create(
        model="gpt-4o-2024-05-13",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "You are a highly skilled AI trained in language comprehension and summarization. I would like you to read the following text and summarize it into a concise abstract paragraph. Aim to retain the most important points, providing a coherent and readable summary that could help a person understand the main points of the discussion without needing to read the entire text. Please avoid unnecessary details or tangential points."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response.choices[0].message.content

def key_points_extraction(transcription):
    response = client.chat.completions.create(
        model="gpt-4o-2024-05-13",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "You are a proficient AI with a specialty in distilling information into key points. Based on the following text, identify and list the main points that were discussed or brought up. These should be the most important ideas, findings, or topics that are crucial to the essence of the discussion. Your goal is to provide a list that someone could read to quickly understand what was talked about."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response.choices[0].message.content

def action_item_extraction(transcription):
    response = client.chat.completions.create(
        model="gpt-4o-2024-05-13",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "You are an AI expert in analyzing conversations and extracting action items. Please review the text and identify any tasks, assignments, or actions that were agreed upon or mentioned as needing to be done. These could be tasks assigned to specific individuals, or general actions that the group has decided to take. Please list these action items clearly and concisely."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response.choices[0].message.content

def sentiment_analysis(transcription):
    response = client.chat.completions.create(
        model="gpt-4o-2024-05-13",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "As an AI with expertise in language and emotion analysis, your task is to analyze the sentiment of the following text. Please consider the overall tone of the discussion, the emotion conveyed by the language used, and the context in which words and phrases are used. Indicate whether the sentiment is generally positive, negative, or neutral, and provide brief explanations for your analysis where possible."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response.choices[0].message.content

def save_as_docx(minutes, filename):
    doc = Document()
    for key, value in minutes.items():
        # Replace underscores with spaces and capitalize each word for the heading
        heading = ' '.join(word.capitalize() for word in key.split('_'))
        doc.add_heading(heading, level=1)
        doc.add_paragraph(value)
        # Add a line break between sections
        doc.add_paragraph()
    doc.save(filename)

audio_file_path = "Amazon-Earnings-Call-Q1-2024-Full-Call-v1.wav"
transcription = transcribe_audio(audio_file_path)
minutes = meeting_minutes(transcription)
print(minutes)

save_as_docx(minutes, 'AMZN_meeting_minutes.docx')
